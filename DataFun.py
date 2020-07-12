import os
import pandas as pd
from datetime import datetime
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE
from docx.enum.table import WD_ALIGN_VERTICAL

Subs =[]
Subs.append('\\New folder\\')


def GetVaildClients(Data):
        #Data = Data.notnull()
        Data = Data.loc[Data['Active'] == 'Y']
        return Data

def GetNames(Data):
    ls = Data[['Client First Name', 'Client Last Name']].apply(lambda x: ''.join(x), axis=1)
    ls = ls.tolist()
    return ls 

def CreateClient(Name):
    
    os.makedirs(Name)

#Big Balla Brand Alert.

def SendForNotes(Notes,Name,Date,Path,Treat,Dig,SessDate):

     
    
    DigCode = MatchCode(Subs[0],'Diagnosis Codes.csv',Dig)
    TreatCode = MatchCode(Subs[0],'Treatment Codes.csv',Treat)
   


    FullerPath = Path  + '\\' + Name+ '\\' +Name+ ".docx"


    print(FullerPath)
    print(Name)
    print(Date)
    print(Notes)
    doc = Document(FullerPath)
    

    


    table = doc.add_table(3,5)
    #table.style = 'Table Grid'
    table.rows[0].cells[0].text = "Session Date"
    table.rows[0].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.rows[1].cells[0].text = SessDate

    table.rows[0].cells[1].text = "Notes Date"
    table.rows[0].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.rows[1].cells[1].text = Date
    
    
    table.rows[0].cells[2].text = "Patient Name"
    table.rows[0].cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.rows[1].cells[2].text = Name

    table.rows[0].cells[3].text = "Dignosis Code"
    table.rows[0].cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.rows[1].cells[3].text = DigCode

    table.rows[0].cells[4].text = "Treatment Code"
    table.rows[0].cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.rows[1].cells[4].text = TreatCode
    
    
    
    
    
    
    table.autofit

    table.cell(2,0).merge(table.cell(2,1))
    table.cell(2,0).merge(table.cell(2,2))
    table.cell(2,0).merge(table.cell(2,3))

    table.cell(2,0).text = Notes

    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    
                    font.name = "Times New Roman"


 
    

    table.alignment = WD_TABLE_ALIGNMENT.CENTER



     #seendate writedate name 
   
    
    doc.save(FullerPath)
    


def FindFile(Subs, File):
    Vars = os.path.abspath('')
    SubFolder = Subs
    MastaFile = File

    return Vars + SubFolder + File
      
           
def OnStart():

    

    PepsData =  pd.DataFrame(pd.read_csv(FindFile(Subs[0], 'MasterList.csv')))
    FileV = FindFile(Subs[0],'')
    PepsDataUpdated = GetVaildClients(PepsData)
    ClientNames = GetNames(PepsDataUpdated)
    now = datetime.now()
    date_time = now.strftime("%m/%d/%Y")

    return(ClientNames,date_time, FileV)

def GetCode(File):
    
    
    Table = pd.DataFrame(pd.read_csv(FindFile(Subs[0], File)))
    ls = Table.iloc[:,0]
    ls = ls.tolist()
    #print(ls)
    return ls 

def MatchCode(Subs, File, Search):
    print(Search)
    Table = pd.DataFrame(pd.read_csv(FindFile(Subs, File))).astype(str)
    #print(Table)
    Val = Table.iloc[:, 0].str.match(Search).loc[lambda x: x==True].index.tolist()
    print(Val)
    Val = Val[0]
    
    Final = Table.iloc[Val,1]
    
    
    #print(Value)
    #StriValue = str(Value)
    return Final


#MatchCode(Subs[0],'Treatment Codes.csv','Individual session_30 minutes with patient')
        
#Modifier to Treatment Code (Teletherapy)		-95
        #ID	Active	Client First Name	Client Last Name	Phone Number	Email	Parent First Name	Parent Last Name


